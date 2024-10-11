import pandas as pd
from docx import Document
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# 读取CSV文件
csv_file = 'SI.csv'  # 替换为你的CSV文件名
df = pd.read_csv(csv_file)

# 创建一个新的Word文档
doc = Document()

# 添加一个表格，并设置行数和列数
table = doc.add_table(rows=df.shape[0] + 1, cols=df.shape[1])

# 设置字体大小（可以根据需求调整此值）
font_size = 10

# 定义一个函数来设置单元格字体大小和统一字体
def set_cell_font(cell, text, font_size):
    # 确保单元格内容的第一个段落有文本
    paragraph = cell.paragraphs[0]
    run = paragraph.runs[0] if paragraph.runs else paragraph.add_run(text)
    run.font.size = Pt(font_size)
    run.font.name = 'Times New Roman'  # 设置统一字体为Times New Roman

# 获取表格的表头行（第一行）
hdr_cells = table.rows[0].cells

# 填充表头数据并设置字体
for i, col_name in enumerate(df.columns):
    hdr_cells[i].text = str(col_name)
    set_cell_font(hdr_cells[i], col_name, font_size)

# 填充表格数据并设置字体
for i, row in df.iterrows():
    row_cells = table.rows[i + 1].cells
    for j, cell_value in enumerate(row):
        row_cells[j].text = str(cell_value)
        set_cell_font(row_cells[j], str(cell_value), font_size)

# 设置表格样式：交替的灰白色行，并且去掉边框
for i, row in enumerate(table.rows):
    if i == 0:
        # 设置表头上下边框
        for cell in row.cells:
            tc = cell._element
            tcPr = tc.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')

            # 设置表头上边框
            top_border = OxmlElement('w:top')
            top_border.set(qn('w:val'), 'single')
            top_border.set(qn('w:sz'), '6')
            top_border.set(qn('w:space'), '0')
            top_border.set(qn('w:color'), '000000')
            tcBorders.append(top_border)

            # 设置表头下边框
            bottom_border = OxmlElement('w:bottom')
            bottom_border.set(qn('w:val'), 'single')
            bottom_border.set(qn('w:sz'), '6')
            bottom_border.set(qn('w:space'), '0')
            bottom_border.set(qn('w:color'), '000000')
            tcBorders.append(bottom_border)

            tcPr.append(tcBorders)
    elif i == len(table.rows) - 1:
        # 设置最后一行的下边框
        for cell in row.cells:
            tc = cell._element
            tcPr = tc.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')

            # 设置最后一行的下边框
            bottom_border = OxmlElement('w:bottom')
            bottom_border.set(qn('w:val'), 'single')
            bottom_border.set(qn('w:sz'), '6')
            bottom_border.set(qn('w:space'), '0')
            bottom_border.set(qn('w:color'), '000000')
            tcBorders.append(bottom_border)

            tcPr.append(tcBorders)
    else:
        # 交替灰白行样式
        if i % 2 == 1:  # 灰色行
            for cell in row.cells:
                tc = cell._element
                tcPr = tc.get_or_add_tcPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:val'), 'clear')
                shd.set(qn('w:fill'), 'D9D9D9')  # 灰色背景
                tcPr.append(shd)

# 设置表格无边框（除表头和尾行的特殊要求外）
for row in table.rows:
    for cell in row.cells:
        tc = cell._element
        tcPr = tc.get_or_add_tcPr()
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)  # 清空边框

# 保存Word文档
doc.save('output_table.docx')

print("Word表格生成完成！")
