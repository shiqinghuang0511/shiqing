import os
import re
from docx import Document
from docx.shared import Pt, Inches

def extract_thermodynamic_data(log_file):
    thermodynamic_data = {}
    
    with open(log_file, 'r') as file:
        lines = file.readlines()

    patterns = {
        "Zero-point correction": r"Zero-point correction=\s+([\d.-]+)",
        "Thermal correction to Energy": r"Thermal correction to Energy=\s+([\d.-]+)",
        "Thermal correction to Enthalpy": r"Thermal correction to Enthalpy=\s+([\d.-]+)",
        "Thermal correction to Gibbs Free Energy": r"Thermal correction to Gibbs Free Energy=\s+([\d.-]+)",
        "Sum of electronic and zero-point Energies": r"Sum of electronic and zero-point Energies=\s+([\d.-]+)",
        "Sum of electronic and thermal Energies": r"Sum of electronic and thermal Energies=\s+([\d.-]+)",
        "Sum of electronic and thermal Enthalpies": r"Sum of electronic and thermal Enthalpies=\s+([\d.-]+)",
        "Sum of electronic and thermal Free Energies": r"Sum of electronic and thermal Free Energies=\s+([\d.-]+)"
    }

    for key, pattern in patterns.items():
        for line in lines:
            match = re.search(pattern, line)
            if match:
                thermodynamic_data[key] = match.group(1)
                break

    return thermodynamic_data

def extract_coordinates(log_file):
    coordinates = []
    atomnos = []
    with open(log_file, 'r') as file:
        lines = file.readlines()

    start_reading = False

    for i, line in enumerate(lines):
        if "Standard orientation" in line:
            start_reading = True
            coordinates = []
            atomnos = []
            for j in range(i+5, len(lines)):  # Skip header and start reading coordinates
                if "---" in lines[j]:
                    break
                tokens = lines[j].split()
                if len(tokens) >= 6:
                    atomnos.append(tokens[1])  # Atom number
                    coordinates.append([float(tokens[3]), float(tokens[4]), float(tokens[5])])

    if not coordinates:
        return None, None  # No coordinates found

    return coordinates, atomnos

def write_to_word(doc, log_file, thermodynamic_data, coordinates, atomnos):
    # Use file name as molecule name
    molecule_name = os.path.basename(log_file).replace('.log', '')

    # Write molecule name
    p = doc.add_paragraph()
    run = p.add_run(molecule_name)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.font.bold = True  # Molecule name in bold
    p.paragraph_format.space_after = Pt(0)  # No space after

    # Write thermodynamic data
    p = doc.add_paragraph('Thermodynamic Data (given by Hartree):')
    p.style.font.name = 'Times New Roman'
    p.style.font.size = Pt(10)
    p.style.font.bold = False
    p.paragraph_format.space_after = Pt(0)  # No space after

    for key, value in thermodynamic_data.items():
        p = doc.add_paragraph(f'{key}: {value}')
        p.style.font.name = 'Times New Roman'
        p.style.font.size = Pt(10)
        p.paragraph_format.space_after = Pt(0)  # No space after

    # Write coordinates
    p = doc.add_paragraph('Optimized Coordinates:')
    p.style.font.name = 'Times New Roman'
    p.style.font.size = Pt(10)
    p.style.font.bold = False
    p.paragraph_format.space_after = Pt(0)  # No space after

    if coordinates and atomnos:
        coord_table = doc.add_table(rows=1, cols=4)
        coord_table.style = 'Table Grid'
        coord_table.autofit = True
        hdr_cells = coord_table.rows[0].cells
        hdr_cells[0].text = 'Atom No.'
        hdr_cells[1].text = 'X (Angstroms)'
        hdr_cells[2].text = 'Y (Angstroms)'
        hdr_cells[3].text = 'Z (Angstroms)'

        for atom, coord in zip(atomnos, coordinates):
            row_cells = coord_table.add_row().cells
            row_cells[0].text = str(atom)
            row_cells[1].text = f'{coord[0]:.6f}'
            row_cells[2].text = f'{coord[1]:.6f}'
            row_cells[3].text = f'{coord[2]:.6f}'

        # Adjust table row spacing
        for row in coord_table.rows:
            for cell in row.cells:
                cell.paragraphs[0].paragraph_format.line_spacing = Pt(10)
                cell.paragraphs[0].paragraph_format.space_after = Pt(0)  # No space after
    else:
        doc.add_paragraph('No optimized coordinates found')

def process_log_files():
    doc = Document()

    # Set document page layout
    section = doc.sections[0]
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    log_files = [f for f in os.listdir('.') if f.endswith('.log')]

    for log_file in log_files:
        thermodynamic_data = extract_thermodynamic_data(log_file)
        coordinates, atomnos = extract_coordinates(log_file)
        write_to_word(doc, log_file, thermodynamic_data, coordinates, atomnos)

    # Save to Word file
    doc.save('cartesian_coordinates.docx')

# Execute script
process_log_files()
